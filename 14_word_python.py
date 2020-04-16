import docx

# Read a word document
d = docx.Document('demo.docx')

# All its paragraphs objects
print(d.paragraphs) 

# One paragraph object
print(d.paragraphs[0])

# Text inside the paragraph
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

# List of run objects inside a paragraph
# A run object starts when there is a change in the style
p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)

# Change the text inside a paragraph
p.runs[2].text = 'italic and underline'
p.runs[2].underline = True

# Save changes into word document
d.save('demo2.docx')

# Changes styles in the word document
p.style = 'Title'
d.save('demo3.docx')

# Creates a word document 
d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo4.docx')

# Add a run(text) to a paragraph
p = d.paragraphs[0]
p.add_run(' This is a new run')
p.runs[1].bold = True # bold
d.save('demo4.docx')

# Get all of the text in a word document
def get_text(word_document):
    doc = docx.Document(word_document) # open it
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

print(get_text('demo.docx'))

# If you want to modify a word document,
# read the old one and create a new one 
# adding the old and new paragraphs to it